
### Improt Modules ###

import os
import numpy as np
import pandas as pd
from pathlib import Path
import pytesseract
from PIL import Image
from PIL import ImageFilter
#from PIL.ExifTags import TAGS 
import re
import math
import unicodedata
import textdistance
import nltk
#from nltk.corpus import stopwords
#from nltk.corpus import stopwords.words('english') as stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from docx import Document
from abc import abstractmethod
from pdf2image import convert_from_path


############################################# Classes on text extraction and text processing #############################################

class extract_text_and_process:

    def __init__(self):
        pass
    
    @abstractmethod
    def extract_text(self, list_paths):
        pass

    def unicode_to_ascii(self, s):
        self.s = s
        return ''.join(c for c in unicodedata.normalize('NFD', self.s) if unicodedata.category(c) != 'Mn')

    def text_cleaning(self, w):
        self.w = w
        self.w = self.unicode_to_ascii(self.w)
        self.w = self.w.lower()                        # Lower casing
        self.w = re.sub(' +', ' ', self.w).strip(' ')  # Remove multiple whitespaces, also leading and trailing whitespaces
        self.w = re.sub(r'[^\w\s]','', self.w)          # Remove special characters and punctuation
        #w = re.sub(r"([0-9])", r" ",w)       # Remove Numerical data
        self.w = re.sub("(.)\\1{2,}", "\\1", self.w)   # Remove duplicate characters
        self.words = self.w.split()                  # Tokenization
        #self.stopwords = nltk.corpus.stopwords.words('english')
        #self.stopwords = stopwords.words('english')
        #self.words = [word for word in self.words if (word not in self.stopwords) and len(word) > 2]
        #self.words = [word for word in self.words if (word not in stopwords) and len(word) > 2]
        self.words = [word for word in self.words if len(word) > 2]
        return " ".join(self.words)

    def process_all_text(self, list_w):
        self.list_w = list_w
        self.clean_strings = []
        self.count = len(self.list_w)
        for i in range(self.count):
            self.clean_text = self.text_cleaning(self.list_w[i])
            self.clean_strings.append(self.clean_text)
        return self.clean_strings



class extract_text_from_image(extract_text_and_process):

    def __init__(self):
        pass
    
    def check_and_adjust_dpi(self, img, size_w=600, size_h=600):
        #size = 7016, 4961
        self.img = img
        self.threshold_w, self.threshold_h = 300, 300
        self.size_w, self.size_h = 600, 600    
        #print("$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$")
        if (self.img.size[0] < self.threshold_w) and (self.img.size[1] < self.threshold_h):
            self.img = self.img.resize((self.size_w, self.size_h))
        return self.img

    def noise_filters(self, img):
        self.img = img
        self.img = self.img.filter(ImageFilter.MinFilter(3)) # Min pix : Be careful in using this !
        return self.img

    def preprocess_image(self, img):
        self.img = img
        
        #print("@@@@@@@@@@@@@@@@@@@@")
        self.img = self.check_and_adjust_dpi(self.img) # adjust dots per image
        #self.img = self.noise_filters(self.img)
        self.img = self.img.convert('L') # binarize
        return self.img
        
    def extract_text(self, list_paths):   
        self.list_paths = list_paths
        self.count = len(self.list_paths)
        self.list_texts = []
        self.path_tesseract = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract

        for i in range(self.count):
            
            if Path(self.list_paths[i]).suffix == 'pdf': # it should be .pdf
                doc = convert_from_path(self.list_paths[i], poppler_path=r'C:\Program Files\poppler-23.08.0\Library\bin')
                self.text = ''
                for page_number, page_data in enumerate(doc):
                    #page_data = self.preprocess_image(page_data)
                    txt = pytesseract.image_to_string(page_data)
                    self.text += txt
            else:
                self.image = Image.open(self.list_paths[i])
                self.image = self.preprocess_image(self.image)
                self.text = pytesseract.image_to_string(self.image)
            
            print("^^^^^^^^^^^^^^^^^^^^^^^^")
            print(i)
            print(self.text)
            self.list_texts.append(self.text)
        
        return self.list_texts
    


class extract_text_from_doc(extract_text_and_process):

    def __init__(self):
        pass

    def extract_text(self, list_paths):   
        self.list_paths = list_paths
        self.texts_list = []
        for i in range(len(self.list_paths)):
            
            self.string_one_file = '' # placeholder !
            
            if Path(self.list_paths[i]).suffix == '.txt':
                with open(self.list_paths[i]) as f:
                    contents = f.read()
                    self.string_one_file = self.string_one_file + contents
                
            elif Path(self.list_paths[i]).suffix == '.docx':
                self.doc = Document(self.list_paths[i])
                for j in range(len(self.doc.paragraphs)):
                    self.string_one_file = self.string_one_file + self.doc.paragraphs[j].text

            self.texts_list.append(self.string_one_file)
             
        return self.texts_list
    


class extract_text_from_table(extract_text_and_process):

    def __init__(self):
        pass

    def extract_text(self, list_paths):
        pass

############################################# Process attachments #############################################

class process_attachments:

    def __init__(self):
        pass

    def detect_file_type(self, file_path):

        self.file_type = None
        self.path = file_path
        if Path(self.path).suffix in ['.csv', '.xls', '.xlsb', '.xlsm', '.xlsx', '.xml', '.ods']:
            self.file_type = "Tabular Data"
        elif Path(self.path).suffix in ['.jpeg', '.jpg', '.png']:
            self.file_type = "Image Data"
        elif Path(self.path).suffix in ['.txt', '.docx']:
            self.file_type = "Text Data"
        else:
            self.file_type = "Junk"
        return self.file_type

    def group_similar_file_types(self, list_paths):

        self.list_table_types, self.list_image_types, self.list_text_types = [], [], []
        self.index_table_types, self.index_image_types, self.index_text_types = [], [], [] 
        self.list_path = list_paths
        
        #for path in list_paths:
        for i in range(len(list_paths)):
            path = list_paths[i]
            self.file_type = self.detect_file_type(path)
            if self.file_type == "Tabular Data":
                self.list_table_types.append(path)
                self.index_table_types.append(str(i+1))
            elif self.file_type == "Image Data":
                self.list_image_types.append(path)
                self.index_image_types.append(str(i+1))
            elif self.file_type == "Text Data":
                self.list_text_types.append(path)
                self.index_text_types.append(str(i+1))
            else:
               pass # Do nothing with that document :: throw it !
        return (self.list_table_types, self.list_image_types, self.list_text_types), (self.index_table_types, self.index_image_types, self.index_text_types)
    

############################################# Plagiarism calculation #############################################
    
class plagiarism_calculation:
    
    def __init__(self):
        pass

    '''
    def __init__(self, ocr_texts, doc_texts, table_texts):
        self.ocr_texts_list = ocr_texts
        self.doc_texts_list = doc_texts
        self.table_texts_list = table_texts
    '''

    def similarity_score(self, list_w, index_types):

        self.list_w = list_w # ['ocr1', 'text2', 'text1', 'ocr2', 'ocr3', 'ocr4]
        self.index_types = index_types # [1, 2, 3, 4, 5, 6]
        self.scores = {} 
        self.count = len(self.list_w)
        for i in range(self.count):
            #key = 'Attach_'+str(i+1)
            key = 'Attach_'+self.index_types[i]
            self.scores[key] = [] # placeholder!
            for j in range(self.count):
                cosine_coef = textdistance.cosine(list_w[i], list_w[j])
                perc_dist = round(cosine_coef*100.0, 2)
                #perc_dist = round((math.pi - math.acos(cosine_coef)) * 100 / math.pi, 2)
                #print(type(cosine_coef), type(perc_dist))
                self.scores[key].append(perc_dist)
    
        return self.scores
    
    #{'Attach_1': [, , , , ,], 'Attach_2': [],  .... }

    
    def similarity_score_all_types(self, ocr_texts, doc_texts, table_texts, index_types):
    
        self.ocr_texts_list = ocr_texts
        self.doc_texts_list = doc_texts
        self.table_texts_list = table_texts
        #self.index_types = index_types
        if ((len(self.ocr_texts_list) !=0) or (len(self.doc_texts_list) !=0)) and (len(self.ocr_texts_list)+len(self.doc_texts_list) >= 2):
            self.ocr_texts_list.extend(self.doc_texts_list) # ['ocr1', 'ocr2', 'ocr3', 'ocr4', 'text2', 'text1']
                                                            #['ocr1', 'text2', 'text1', 'ocr2', 'ocr3', 'ocr4] # POSTMAN
            self.index_types = list(index_types[1])
            self.index_types.extend(list(index_types[2]))   # ['1', '4', '5', '6', '2', '3']
            print("MUSAWER:", self.index_types)

            temp_list = []

            #for index in self.index_types:
                #x = int(index) # 1
                #temp_list.append(self.ocr_texts_list[x-1])


            for i in range(len(self.index_types)):
                #indx = self.index_types[i-1]
                index = str(i+1)
                true_index = self.index_types.index(index)
                temp_list.append(self.ocr_texts_list[true_index])

            
            self.index_types.sort() # [1, 2, 3, 4, 5, 6]
            self.temp_list = temp_list

            self.score_matrix = self.similarity_score(self.temp_list, self.index_types)
        
        else:
            self.score_matrix = {"Attach_None": "NA"}

        #self.score_matrix = self.similarity_score(self.ocr_text_list)

        '''
        if len(self.list_table_texts) != 0:
            # concate all the rows into a single dataframe, compute similarity among the rows
            # and then compare each row with the text if there is any
            self.combined_df = pd.DataFrame()

            for table in self.list_table_texts:
                df = pd.read_csv(table)
                self.combined_df = pd.concat([self.combined_df, df])
                #match_rows(combined_df)
        '''
        return self.score_matrix
    

    def filter_top_sim_score(self, score_matrix):
        
        self.doc_names = list(score_matrix.keys()) # [Attach_1, Attach_2, Attach_3, Attach_4, Attach_5, Attach_6]
        self.primary_output = []
        
        for key in score_matrix.keys():

            temp_arr = np.array(score_matrix[key])
            highest = np.partition(temp_arr.flatten(), -2)[-2]
            index = np.where(temp_arr == highest)[0][0]
            temp_list = [key, self.doc_names[index], highest]
            self.primary_output.append(temp_list)

        self.final_output = {}
        self.final_output['primary_output'] = self.primary_output
        self.final_output['secondary_output'] = score_matrix

        return self.final_output


    def filter_matrix(self, score_matrix):
        
        df = pd.DataFrame(columns=list(score_matrix.keys()))

        for key in score_matrix:
            df[key] = score_matrix[key]       
        df = df.T
        df.columns = list(score_matrix.keys())
        #print(df)
        df = df.where(np.triu(np.ones(df.shape)).astype(bool))
        #print(df)
        df = df.stack().reset_index()
        df.columns = ['doc_1', 'doc_2', 'perc_sim']
        #print(df)
        if len(df[df['doc_1'] == df['doc_2']]) < len(df):
            #df = df.drop_duplicates(subset=['Column2'], keep=False)
            df.drop(df[df['doc_1'] == df['doc_2']].index, inplace = True)
            df = df.sort_values(by='perc_sim', ascending=False)
            df.reset_index(drop=True, inplace=True)
        df = df.head(5)
        #print("***************************************")
        #print(df)
        self.primary_output = df.values.tolist()
        #print("---------------------------------------")
        #print(df)
        self.final_output = {}
        self.final_output['primary_output'] = self.primary_output
        self.final_output['secondary_output'] = score_matrix

        return self.final_output

        

        








