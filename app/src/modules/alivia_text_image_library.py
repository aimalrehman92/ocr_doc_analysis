

import os
import numpy as np
import pandas as pd
from pathlib import Path
import pytesseract
from PIL import Image
from PIL import ImageFilter, ImageOps
#from PIL.ExifTags import TAGS 
import re
import math
import unicodedata
import textdistance
import nltk
import cv2
#from nltk.corpus import stopwords
#from nltk.corpus import stopwords.words('english') as stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from docx import Document
from abc import abstractmethod
from pdf2image import convert_from_path

from datetime import datetime
import PyPDF2

pd.options.mode.chained_assignment = None  # or 'warn'

############################################# Classes on text extraction and text processing #############################################


class ExtractTextAndProcess:

    def __init__(self):
        pass
    
    @abstractmethod
    def extract_text(self, list_paths):
        pass


    def unicode_to_ascii(self, s):
        
        self.s = s
        return ''.join(c for c in unicodedata.normalize('NFD', self.s) if unicodedata.category(c) != 'Mn')


    def text_cleaning(self, w):
        
        self.w = w
        self.w = self.unicode_to_ascii(self.w)
        self.w = self.w.lower()                        # Lower casing
        self.w = re.sub(' +', ' ', self.w).strip(' ')  # Remove multiple whitespaces, also leading and trailing whitespaces
        self.w = re.sub(r'[^\w\s]','', self.w)         # Remove special characters and punctuation
        #self.w = re.sub(r"([0-9])", r" ", self.w)                # Remove Numerical data
        #self.w = re.sub("(.)\\1{2,}", "\\1", self.w)   # Remove duplicate characters
        self.words = self.w.split()                  # Tokenization
        #self.stopwords = nltk.corpus.stopwords.words('english')
        #self.stopwords = stopwords.words('english')
        #self.words = [word for word in self.words if (word not in self.stopwords) and len(word) > 2]
        #self.words = [word for word in self.words if (word not in stopwords) and len(word) > 2]
        self.words = [word for word in self.words if len(word) > 2]
        return " ".join(self.words)


    def process_all_text(self, list_w):
        
        self.list_w = list_w
        self.clean_strings = []
        self.count = len(self.list_w)
        for i in range(self.count):
            self.clean_text = self.text_cleaning(self.list_w[i])
            self.clean_strings.append(self.clean_text)
        return self.clean_strings
    

    def process_single_string(self, single_string):
        
        #clean_text = self.text_cleaning(single_string)
        clean_text = single_string.split()
        return clean_text


    def remove_special_characters_word(self, words_list):

        if len(words_list) > 0:
            words_list = [word for word in words_list if not re.match(r'^\W+$', word)]
        return words_list


class ExtractImageText(ExtractTextAndProcess):

    def __init__(self, flags_dict):

        self.color_to_greyscale_flag = flags_dict.get('color_to_greyscale', True)
        self.adjust_image_dpi_flag = flags_dict.get('adjust_dpi', True)
        self.noise_filters_flag = flags_dict.get('noise_filters', False)
        self.binarize_image_flag = flags_dict.get('binarize_image', False)
        self.adjust_image_size_flag = flags_dict.get('adjust_image_size', True)
        self.resize_to_A4_flag = flags_dict.get('resize_to_A4', True)
        self.equalizeHist_flag = False

        self.path_tesseract = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        self.poppler_path = os.getcwd() + r'\app\src\modules\poppler-23.08.0\Library\bin'

    
    def adjust_image_dpi(self, image, scale_h=2, scale_w=2):
        
        if self.check_image_type(image) == "Pillow image":
            
            if ('dpi' in image.info) and (image.info['dpi'][0] < 2*300 or image.info['dpi'][1] < 2*300):
                image.info['dpi'] = (20*300, 20*300) 
             
        if self.check_image_type(image) == "Numpy array":
            
            threshold_w, threshold_h = 300, 300
            h, w = image.shape[0], image.shape[1]
            
            if (w < threshold_w) and (h < threshold_h):
                scale_h, scale_w = max(2, scale_h), max(2, scale_w)
                image = cv2.resize(image, (h*scale_h, w*scale_w))
            else:
                image = cv2.resize(image, (h*scale_h, w*scale_w))

        return image
    

    def resize_to_A4(self, image):

        if self.check_image_type(image) == "Numpy array":
            image = cv2.resize(image, (2480, 3508)) # A4 resolution
        
        if self.check_image_type(image) == "Pillow image":
            pass

        return image


    def noise_filters(self, image):
        
        image = image.filter(ImageFilter.MinFilter(3)) # Min pix : Be careful in using this !
        return image
    
 
    def binarize_image(self, image, threshold=120):

        if self.check_image_type(image) == "Pillow image":
            image = image.point(lambda p: 0 if p < threshold else 255)
            
        if self.check_image_type(image) == "Numpy array":
            _, image = cv2.threshold(image, threshold, 255, cv2.THRESH_BINARY)

        return image


    def check_image_type(self, image):

        if isinstance(image, Image.Image):
            image_type = "Pillow image"
        elif isinstance(image, np.ndarray):
            image_type =  "Numpy array"
        else:
            image_type = "Not a valid image"
        
        return image_type
    

    def color_to_greyscale(self, image):

        if self.check_image_type(image) == "Numpy array":
            if (len(image.shape) >= 3) and (image.shape[2] == 3):
                image = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY) 
        
        elif self.check_image_type(image) == "Pillow image":
            if image.mode != 'L':
                image = ImageOps.grayscale(image) # preferred method
                #image = image.convert("L") # less-flexible method

        return image
    
    
    def equalize_text_hist(self, image):
        
        if self.check_image_type(image) == "Numpy array":
            image = cv2.equalizeHist(image)

        elif self.check_image_type(image) == "Pillow image":
            image = ImageOps.equalize(image)

        return image
    
            
    def is_black_background_histogram(self, image_path):
        
        image = cv2.imread(image_path)

        if image is None:
            raise ValueError("Could not load the image.")

        gray_image = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY) # Lock it on RGB and the always read it in RGB REMEMBER !
        hist = cv2.calcHist([gray_image], [0], None, [256], [0, 256])
        black_pixel_sum = np.sum(hist[:20])
        white_pixel_sum = np.sum(hist[220:256])

        if black_pixel_sum > white_pixel_sum:
            return True  # Likely a black background with white text
        else:
            return False  # Probably not black background with white text


    def adjust_image_size(self, image, adjusted_resolution):

        # This function will increase the dimensions of image (if low) by appending blank spaces around it

        if self.check_image_type(image) == "Pillow image":
            
            if (adjusted_resolution[0] > image.width) and (adjusted_resolution[1] > image.height):
                
                new_width, new_height = adjusted_resolution[0], adjusted_resolution[1]
           
                empty_image = Image.new("RGB", (new_width, new_height), (255, 255, 255))  # White background
                empty_image_copy = Image.new("RGB", (new_width, new_height), (255, 255, 255))  # White background

                paste_x = (new_width - image.width) // 2
                paste_y = (new_height - image.height) // 2

                empty_image.paste(image, (paste_x, paste_y))

                resized_image = empty_image
            
            else:

                resized_image = image
                empty_image_copy = image


        elif self.check_image_type(image) == "Numpy array":
            
            if (adjusted_resolution[0] > image.shape[0]) and (adjusted_resolution[1] > image.shape[1]):
                
                new_width = adjusted_resolution[1]
                new_height = adjusted_resolution[0]
            
                empty_image = np.ones((new_width, new_height), dtype=np.uint8) * 255  # Create a white image
                empty_image_copy = np.ones((new_height, new_width), dtype=np.uint8) * 255

                x_offset = (empty_image.shape[1] - image.shape[1]) // 2
                y_offset = (empty_image.shape[0] - image.shape[0]) // 2

                x_end = x_offset + image.shape[1]
                y_end = y_offset + image.shape[0]

                # Paste the small image onto the empty image at the center
                empty_image[y_offset:y_end, x_offset:x_end] = image

                resized_image = empty_image

            else:
                resized_image = image
                empty_image_copy = image

        return resized_image, empty_image_copy


    def preprocess_image(self, image, adjusted_resolution=(1920, 1080)):
        
        if self.color_to_greyscale_flag:
            image = self.color_to_greyscale(image)
            
        if self.adjust_image_dpi_flag:
            image = self.adjust_image_dpi(image) # adjust dots per image
            
        if self.noise_filters_flag:
            image = self.noise_filters(image)
                
        if self.binarize_image_flag:
            image = self.binarize_image(image)
            
        if self.equalizeHist_flag:
            image = self.equalize_text_hist(image)

        return image


    def postprocess_images(self, images_list, adjusted_resolution=(1920, 1080)):

        for j in range(len(images_list)):

            image = images_list[j]
        
            if self.adjust_image_size_flag: # radjust the dimensions
                image, _ = self.adjust_image_size(image, adjusted_resolution)

            if self.resize_to_A4_flag:
                image = self.resize_to_A4(image)

            images_list[j] = image
    
        return images_list
        

    def extract_text(self, file_path):   
        
        # This function will return the entire text from one document (PDF, jpg or another image) in the form of a string

        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract

        texts = ''

        if Path(file_path).suffix == '.pdf':
            doc = convert_from_path(file_path, poppler_path=self.poppler_path)
            for page_number, page_data in enumerate(doc):
                page_data = self.preprocess_image(page_data) # we will put self.flags inside here
                texts += pytesseract.image_to_string(page_data)
                #texts = self.postprocess_text(texts) # we will put self.flags inside here

        else:
            image = Image.open(file_path) # Pillow image
            image = self.preprocess_image(image) # we will put self.flags inside here
            texts = pytesseract.image_to_string(image)
            #texts = self.postprocess_text(texts) # we will put self.flags inside here
        
        return texts
    
    
    def extract_text_with_coordinates(self, path):   
        
        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract 
        
        meta_data_df = pd.DataFrame()
        list_page_images = []
        
        if Path(path).suffix == '.pdf':
            
            doc = convert_from_path(path, poppler_path=self.poppler_path)
            
            for page_number, page_data in enumerate(doc):
                
                frame = pd.DataFrame(pytesseract.image_to_data(page_data, output_type=pytesseract.Output.DICT))
                frame['page_num'] = page_number + 1
                meta_data_df = pd.concat([meta_data_df, frame], ignore_index=True) # meta data
                #image = cv2.cvtColor(np.array(page_data), cv2.COLOR_BGR2RGB)
                image = Image.fromarray(np.array(page_data).astype('uint8'), 'RGB') # image in Numpy array so OpenCV
                image = self.preprocess_image(image)
                list_page_images.append(image)

        else:
            
            image = np.array(Image.open(path).convert('RGB')) # image data in numpy array
            image = self.preprocess_image(image)
            meta_data_df = pd.DataFrame(pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)) # meta data
            list_page_images.append(image) # list of only one image !

        return meta_data_df, list_page_images



class ExtractDocumentText(ExtractTextAndProcess):

    def __init__(self):
        pass

    def extract_text(self, list_paths):   
        
        self.list_paths = list_paths
       
        self.string_one_file = '' # placeholder !
        if Path(self.list_paths).suffix == '.txt':
            with open(self.list_paths) as f:
                contents = f.read()
                self.string_one_file = self.string_one_file + contents
        
        elif Path(self.list_paths).suffix == '.docx':
            self.doc = Document(self.list_paths)
            for j in range(len(self.doc.paragraphs)):
                self.string_one_file = self.string_one_file + self.doc.paragraphs[j].text
          
        self.texts_list = self.string_one_file
 
        return self.texts_list
    

class ExtractTableText(ExtractTextAndProcess):

    def __init__(self):
        pass

    def extract_text(self, list_paths):
        pass



class ReturnImageData:

    def __init__(self):
        #super(PreProcessing,self).__init__()
        pass

    def null_image(self, dimensions):
        image = np.zeros(dimensions)
        return image


    def cv_image(self, path):
        image = cv2.imread(path)
        return image
    

    def split_and_repopulate_string_list(self, string_list):
        new_list = []
        for s in string_list:
            if " " in s: #or not s.isalpha():
                # Split the string by spaces and add resulting words to the new list
                words = s.split()
                new_list.extend(words)
            else:
                new_list.append(s)
        return new_list


    def highlight_text_on_image(self, ocr_meta_data, common_words, ocr_images, img_num):
        
        no_of_pages = len(ocr_images)
        common_words =  self.split_and_repopulate_string_list(common_words)

        margin = 6
        for ii in range(no_of_pages): # loop over the pages within a document
            
            ocr_image = ocr_images[ii]  # a page
            
            ocr_image = np.array(ocr_image)

            if len(ocr_image.shape) < 3:
                ocr_image = np.repeat(ocr_image[:, :, np.newaxis], 3, axis=2) # ensure the number of channels is 3 even if its greyscale

            # modfiy the datfarame containing meta data to group neighboring words that are plagiarised
                
            meta_data = ocr_meta_data[ocr_meta_data['page_num'] == ii+1].copy()
            meta_data = meta_data[meta_data['text'].isin(common_words)]
            meta_data['right'] = meta_data['left'] + meta_data['width']
            meta_data['bottom'] = meta_data['top'] + meta_data['height']
            df_sample = meta_data[['text', 'page_num', 'block_num', 'par_num', 'line_num', 'word_num', 'left', 'top', 'right', 'bottom']].copy()
            df_sample['delta_page_num'] = df_sample['page_num'] - df_sample['page_num'].shift(+1)
            df_sample['delta_block_num'] = df_sample['block_num'] - df_sample['block_num'].shift(+1)
            df_sample['delta_par_num'] = df_sample['par_num'] - df_sample['par_num'].shift(+1)
            df_sample['delta_line_num'] = df_sample['line_num'] - df_sample['line_num'].shift(+1)
            df_sample['delta_word_num'] = df_sample['word_num'] - df_sample['word_num'].shift(+1)

            df_sample['delta_page_num'] = df_sample['delta_page_num'].fillna(0.0)
            df_sample['delta_block_num'] = df_sample['delta_block_num'].fillna(0.0)
            df_sample['delta_par_num'] = df_sample['delta_par_num'].fillna(0.0)
            df_sample['delta_line_num'] = df_sample['delta_line_num'].fillna(0.0)
            df_sample['delta_word_num'] = df_sample['delta_word_num'].fillna(1.0)

            blob = 1
            df_sample['blob'] = blob

            for k in range(len(df_sample)):
                if k > 0:
                    if df_sample['delta_word_num'].iloc[k] != 1:
                        blob += 1
        
                df_sample['blob'].iloc[k] = blob

            df_toy_1 = df_sample[(df_sample['delta_page_num']==0)].groupby(['page_num', 'block_num', 'par_num', 'line_num', 'blob'])['left'].min()
            df_toy_1 = df_toy_1.rename('left')

            df_toy_2 = df_sample[(df_sample['delta_page_num']==0)].groupby(['page_num', 'block_num', 'par_num', 'line_num', 'blob'])['right'].max()
            df_toy_2 = df_toy_2.rename('right')

            df_toy_3 = df_sample[(df_sample['delta_page_num']==0)].groupby(['page_num', 'block_num', 'par_num', 'line_num', 'blob'])['top'].min()
            df_toy_3 = df_toy_3.rename('top')

            df_toy_4 = df_sample[(df_sample['delta_page_num']==0)].groupby(['page_num', 'block_num', 'par_num', 'line_num', 'blob'])['bottom'].max()
            df_toy_4 = df_toy_4.rename('bottom')

            df_toy = pd.concat([df_toy_1, df_toy_2, df_toy_3, df_toy_4], axis=1)
            del(df_toy_1, df_toy_2, df_toy_3, df_toy_4)

            # Now highlight those (group of) words by iterating over them
            for _, row in df_toy.iterrows():
                
                left, top, right, bottom = row[['left', 'top', 'right', 'bottom']].values
                sub_img = ocr_image[top : bottom, left : right]
                highlighter = np.array([255, 225, 20], dtype=np.uint8) # yellow shade
                white_rect = np.ones(sub_img.shape, dtype=np.uint8)*highlighter

                alpha = 0.8
                beta = 1-alpha
                
                blended = cv2.addWeighted(sub_img, alpha, white_rect, beta, 1)
                ocr_image[top : bottom, left : right] = blended
        
            #ocr_image = cv2.resize(ocr_image, (2480, 3508)) # A4 resolution

            ocr_images[ii] = ocr_image        

        return ocr_images
    

    def return_text_from_doc(self, common_words, i):
        
        file_path = f"plagiarised_doc_{i}.txt"

        f = open(file_path, "w", encoding='utf-8')        
        for word in common_words:
            f.write(word)
            f.write(" ")

        f.close()

        return file_path
    

    def create_outline(self, pdf_file, document_text, document_meta_data, text_to_mark, keys_found, output_pdf_file):
        
        pdf = PyPDF2.PdfReader(open(pdf_file, "rb"))
        pdf_writer = PyPDF2.PdfWriter()

        for page_num in range(len(pdf.pages)):
            
            page = pdf.pages[page_num]
            pdf_writer.add_page(page)

            page_text = document_meta_data[document_meta_data['page_num'] == page_num+1]['text'].tolist()
            page_text = " ".join(page_text)

            for i in range(len(text_to_mark)):

                text_string = text_to_mark[i]

                if text_string in page_text:
                    bookmark= pdf_writer.add_outline_item(keys_found[i], page_num)

        with open(output_pdf_file, "wb") as output_pdf:
            pdf_writer.write(output_pdf)
        
        output_pdf.close()

        return None
