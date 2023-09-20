

import os
import numpy as np
import pandas as pd
from pathlib import Path
import pytesseract
from PIL import Image
from PIL import ImageFilter
#from PIL.ExifTags import TAGS 
import re
import math
import unicodedata
import textdistance
import nltk
#from nltk.corpus import stopwords
#from nltk.corpus import stopwords.words('english') as stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from docx import Document
from abc import abstractmethod
from pdf2image import convert_from_path

from datetime import datetime

############################################# Classes on text extraction and text processing #############################################

class extract_text_and_process:

    def __init__(self):
        pass
    
    @abstractmethod
    def extract_text(self, list_paths):
        pass

    def unicode_to_ascii(self, s):
        
        self.s = s
        return ''.join(c for c in unicodedata.normalize('NFD', self.s) if unicodedata.category(c) != 'Mn')

    def text_cleaning(self, w):
        
        self.w = w
        self.w = self.unicode_to_ascii(self.w)
        self.w = self.w.lower()                        # Lower casing
        self.w = re.sub(' +', ' ', self.w).strip(' ')  # Remove multiple whitespaces, also leading and trailing whitespaces
        self.w = re.sub(r'[^\w\s]','', self.w)         # Remove special characters and punctuation
        #self.w = re.sub(r"([0-9])", r" ", self.w)                # Remove Numerical data
        #self.w = re.sub("(.)\\1{2,}", "\\1", self.w)   # Remove duplicate characters
        self.words = self.w.split()                  # Tokenization
        #self.stopwords = nltk.corpus.stopwords.words('english')
        #self.stopwords = stopwords.words('english')
        #self.words = [word for word in self.words if (word not in self.stopwords) and len(word) > 2]
        #self.words = [word for word in self.words if (word not in stopwords) and len(word) > 2]
        self.words = [word for word in self.words if len(word) > 2]
        return " ".join(self.words)

    def process_all_text(self, list_w):
        
        self.list_w = list_w
        self.clean_strings = []
        self.count = len(self.list_w)
        for i in range(self.count):
            self.clean_text = self.text_cleaning(self.list_w[i])
            self.clean_strings.append(self.clean_text)
        return self.clean_strings
    

    def process_single_string(self, list_w):
        
        self.clean_text = self.text_cleaning(list_w)
        return self.clean_text


class extract_text_from_image(extract_text_and_process):

    def __init__(self):
        
        self.path_tesseract = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        self.poppler_path=r'C:\Program Files\poppler-23.08.0\Library\bin'


    def check_and_adjust_dpi(self, img, size_w=600, size_h=600):
        
        #size = 7016, 4961
        self.img = img
        self.threshold_w, self.threshold_h = 300, 300
        self.size_w, self.size_h = 600, 600    
        if (self.img.size[0] < self.threshold_w) and (self.img.size[1] < self.threshold_h):
            self.img = self.img.resize((self.size_w, self.size_h))
        return self.img


    def noise_filters(self, img):
        
        self.img = img
        self.img = self.img.filter(ImageFilter.MinFilter(3)) # Min pix : Be careful in using this !
        return self.img


    def preprocess_image(self, img):
        
        self.img = img
        self.img = self.check_and_adjust_dpi(self.img) # adjust dots per image
        #self.img = self.noise_filters(self.img)
        self.img = self.img.convert('L') # binarize
        return self.img


    def extract_text(self, list_paths):   
        
        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract 
        self.list_texts = [] 
        if Path(list_paths).suffix == '.pdf':
            doc = convert_from_path(list_paths, poppler_path=self.poppler_path)
            texts = ''
            for page_number, page_data in enumerate(doc):
                page_data = self.preprocess_image(page_data)
                texts += pytesseract.image_to_string(page_data)
        else:
            img = Image.open(list_paths)
            img = self.preprocess_image(img)
            texts = pytesseract.image_to_string(img)
            #texts = pytesseract.image_to_data(img)
        self.list_texts = texts
        return self.list_texts
    

class extract_text_from_doc(extract_text_and_process):

    def __init__(self):
        pass

    def extract_text(self, list_paths):   
        self.list_paths = list_paths
        self.texts_list = []
        for i in range(len(self.list_paths)):
            self.string_one_file = '' # placeholder !
            if Path(self.list_paths[i]).suffix == '.txt':
                with open(self.list_paths[i]) as f:
                    contents = f.read()
                    self.string_one_file = self.string_one_file + contents
            elif Path(self.list_paths[i]).suffix == '.docx':
                self.doc = Document(self.list_paths[i])
                for j in range(len(self.doc.paragraphs)):
                    self.string_one_file = self.string_one_file + self.doc.paragraphs[j].text
            self.texts_list.append(self.string_one_file)
             
        return self.texts_list
    


class extract_text_from_table(extract_text_and_process):

    def __init__(self):
        pass

    def extract_text(self, list_paths):
        pass


if __name__ == '__main__':
    
    list_paths = [
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Pres 1.jpg",
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Pres 2.jpg",
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Prescription 1.jpg"
    ]

    extract_from_image = extract_text_from_image()

    t0 = datetime.now()
    list_texts = extract_from_image.extract_text(list_paths)
    print("time taken: ", datetime.now()-t0)

    print("Length:", len(list_texts))
    print(list_texts)

