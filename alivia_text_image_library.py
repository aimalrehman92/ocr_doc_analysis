

import os
import numpy as np
import pandas as pd
from pathlib import Path
import pytesseract
from PIL import Image
from PIL import ImageFilter
#from PIL.ExifTags import TAGS 
import re
import math
import unicodedata
import textdistance
import nltk
import cv2
#from nltk.corpus import stopwords
#from nltk.corpus import stopwords.words('english') as stopwords
from nltk.tokenize import word_tokenize,sent_tokenize
from docx import Document
from abc import abstractmethod
from pdf2image import convert_from_path

from datetime import datetime

############################################# Classes on text extraction and text processing #############################################

class ExtractTextAndProcess:

    def __init__(self):
        pass
    
    @abstractmethod
    def extract_text(self, list_paths):
        pass

    def unicode_to_ascii(self, s):
        
        self.s = s
        return ''.join(c for c in unicodedata.normalize('NFD', self.s) if unicodedata.category(c) != 'Mn')

    def text_cleaning(self, w):
        
        self.w = w
        self.w = self.unicode_to_ascii(self.w)
        self.w = self.w.lower()                        # Lower casing
        self.w = re.sub(' +', ' ', self.w).strip(' ')  # Remove multiple whitespaces, also leading and trailing whitespaces
        self.w = re.sub(r'[^\w\s]','', self.w)         # Remove special characters and punctuation
        #self.w = re.sub(r"([0-9])", r" ", self.w)                # Remove Numerical data
        #self.w = re.sub("(.)\\1{2,}", "\\1", self.w)   # Remove duplicate characters
        self.words = self.w.split()                  # Tokenization
        #self.stopwords = nltk.corpus.stopwords.words('english')
        #self.stopwords = stopwords.words('english')
        #self.words = [word for word in self.words if (word not in self.stopwords) and len(word) > 2]
        #self.words = [word for word in self.words if (word not in stopwords) and len(word) > 2]
        self.words = [word for word in self.words if len(word) > 2]
        return " ".join(self.words)

    def process_all_text(self, list_w):
        
        self.list_w = list_w
        self.clean_strings = []
        self.count = len(self.list_w)
        for i in range(self.count):
            self.clean_text = self.text_cleaning(self.list_w[i])
            self.clean_strings.append(self.clean_text)
        return self.clean_strings
    

    def process_single_string(self, single_string):
        
        #clean_text = self.text_cleaning(single_string)
        clean_text = single_string.split()
        return clean_text


class ExtractImageText(ExtractTextAndProcess):

    def __init__(self):
        
        self.path_tesseract = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        #self.poppler_path = r'C:\Program Files\poppler-23.08.0\Library\bin'

        #self.path_tesseract = os.getcwd() + r'\Tesseract-OCR\tesseract-5.3.1\tesseract.exe'
        self.poppler_path = os.getcwd() + r'\poppler-23.08.0\Library\bin'


    def check_and_adjust_dpi(self, img, size_w=600, size_h=600):
        
        #size = 7016, 4961
        self.img = img
        self.threshold_w, self.threshold_h = 300, 300
        self.size_w, self.size_h = 600, 600    
        if (self.img.size[0] < self.threshold_w) and (self.img.size[1] < self.threshold_h):
            self.img = self.img.resize((self.size_w, self.size_h))
        return self.img


    def noise_filters(self, img):
        
        self.img = img
        self.img = self.img.filter(ImageFilter.MinFilter(3)) # Min pix : Be careful in using this !
        return self.img


    def preprocess_image(self, img):
        
        self.img = img
        self.img = self.check_and_adjust_dpi(self.img) # adjust dots per image
        #self.img = self.noise_filters(self.img)
        self.img = self.img.convert('L') # binarize
        return self.img


    def extract_text(self, list_paths):   
        
        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract 
        self.list_texts = [] 
        if Path(list_paths).suffix == '.pdf':
            doc = convert_from_path(list_paths, poppler_path=self.poppler_path)
            texts = ''
            for page_number, page_data in enumerate(doc):
                page_data = self.preprocess_image(page_data)
                texts += pytesseract.image_to_string(page_data)

        else:
            img = Image.open(list_paths)
            img = self.preprocess_image(img)
            texts = pytesseract.image_to_string(img)
        
        self.list_texts = texts
        return self.list_texts
    
    
    def extract_text_with_coordinates(self, path):   
        
        pytesseract.pytesseract.tesseract_cmd = self.path_tesseract 
        
        meta_data_df = pd.DataFrame()
        list_page_images = []
        
        if Path(path).suffix == '.pdf':
            
            doc = convert_from_path(path, poppler_path=self.poppler_path)
            
            for page_number, page_data in enumerate(doc):
                
                frame = pd.DataFrame(pytesseract.image_to_data(page_data, output_type=pytesseract.Output.DICT))
                frame['page_num'] = page_number + 1
                meta_data_df = pd.concat([meta_data_df, frame], ignore_index=True)
                #image = cv2.cvtColor(np.array(page_data), cv2.COLOR_BGR2RGB)
                image = Image.fromarray(np.array(page_data).astype('uint8'), 'RGB')
                list_page_images.append(image)

        else:
            
            image = np.asarray(Image.open(path).convert('RGB'))
            #image = self.preprocess_image(image)
            meta_data_df = pd.DataFrame(pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT))
            list_page_images.append(image)

        #meta_data_df = meta_data_df[meta_data_df['conf'] > 0]
            
        return meta_data_df, list_page_images
    

class ExtractDocumentText(ExtractTextAndProcess):

    def __init__(self):
        pass

    def extract_text(self, list_paths):   
        
        self.list_paths = list_paths
       
        self.string_one_file = '' # placeholder !
        if Path(self.list_paths).suffix == '.txt':
            with open(self.list_paths) as f:
                contents = f.read()
                self.string_one_file = self.string_one_file + contents
        
        elif Path(self.list_paths).suffix == '.docx':
            self.doc = Document(self.list_paths)
            for j in range(len(self.doc.paragraphs)):
                self.string_one_file = self.string_one_file + self.doc.paragraphs[j].text
          
        self.texts_list = self.string_one_file
 
        return self.texts_list
    

class ExtractTableText(ExtractTextAndProcess):

    def __init__(self):
        pass

    def extract_text(self, list_paths):
        pass


class ReturnImageData:

    def __init__(self):
        #super(PreProcessing,self).__init__()
        pass

    def null_image(self, dimensions):
        image = np.zeros(dimensions)
        return image

    def cv_image(self, path):
        image = cv2.imread(path)
        return image
    
    def highlight_text_on_image(self, ocr_meta_data, common_words, ocr_images, img_num):
        
        no_of_pages = len(ocr_images)
        for ii in range(no_of_pages): # loop over the pages within a document
            ocr_image = ocr_images[ii]  # a page
            ocr_image = np.array(ocr_image)
            meta_data = ocr_meta_data[ocr_meta_data['page_num'] == ii+1].copy()
            for text in common_words:
                retrieved_info = meta_data.loc[meta_data['text'] == text, ['left', 'top', 'width', 'height']]
                if len(retrieved_info) > 0:
                    useful_coord = meta_data.loc[meta_data['text'] == text, ['left', 'top', 'width', 'height']].reset_index(drop=True).values[0]
                    x, y, w, h = useful_coord[0], useful_coord[1], useful_coord[2], useful_coord[3]
                    #sub_img = ocr_image[y : y + h, x : x + w]
                    #white_rect = np.ones(sub_img.shape, dtype=np.uint8) * 255
                    #res = cv2.addWeighted(sub_img, 0.2, white_rect, 0.5, 0)
                    #ocr_image[y : y + h, x : x + w] = res
                    cv2.rectangle(ocr_image, (x, y), (x + w, y + h), (0, 0, 255), 2)
            
            ocr_image = cv2.resize(ocr_image, (700, 900))
            ocr_images[ii] = ocr_image        
        # store the PDF or IMAGE file and return its path
        #image_path = f"plagiarised_img_{img_num}.pdf"
        #cv2.imwrite(image_path, ocr_images)
        return ocr_images
    
    
    def highlight_text_on_image_OLD(self, meta_data, common_words, bkg_image, img_num):
        
        #meta_data = pd.DataFrame.from_dict(meta_data)

        for text in common_words:
            
            useful_coord = meta_data.loc[meta_data['text'] == text, ['left', 'top', 'width', 'height']].values[0]
            x, y, w, h = useful_coord[0], useful_coord[1], useful_coord[2], useful_coord[3]
            sub_img = bkg_image[y : y + h, x : x + w]
            white_rect = np.ones(sub_img.shape, dtype=np.uint8) * 255
            res = cv2.addWeighted(sub_img, 0.2, white_rect, 0.5, 0)
            bkg_image[y : y + h, x : x + w] = res
            #cv2.rectangle(bkg_image, (x, y), (x + w, y + h), (255, 255, 0), -1)

        bkg_image = cv2.resize(bkg_image, (800, 1128))
        image_path = f"plagiarised_img_{img_num}.jpg"
        cv2.imwrite(image_path, bkg_image)
        
        return bkg_image, image_path


    def return_text_from_doc(self, common_words, i):
        
        file_path = f"plagiarised_doc_{i}.txt"

        f = open(file_path, "w", encoding='utf-8')        
        for word in common_words:
            f.write(word)
            f.write(" ")

        f.close()

        return file_path
        

if __name__ == '__main__':
    
    list_paths = [
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Pres 1.jpg",
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Pres 2.jpg",
    "C:\\Users\\MuhammadAimalRehman\\Documents\\OCR_Project\\PyTesseract_Demo_01\\pytesser_demo\\Data_for_SimilarityDetection\\OCR_Tahira\\Prescription 1.jpg"
    ]

    extract_from_image = ExtractImageText()

    t0 = datetime.now()
    list_texts = extract_from_image.extract_text(list_paths)
    print("time taken: ", datetime.now()-t0)

    print("Length:", len(list_texts))
    print(list_texts)

